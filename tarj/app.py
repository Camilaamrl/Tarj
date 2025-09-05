from flask import Flask, render_template, request, send_file, jsonify, session
import re
import spacy
from spacy.pipeline import EntityRuler
import tempfile
import os
import fitz
from docx import Document
import io
import uuid
import base64
from docx.shared import RGBColor
import json
from fuzzywuzzy import fuzz
from pdf2image import convert_from_bytes
import pytesseract
from pyzbar.pyzbar import decode
from PIL import Image, ImageDraw

# Lista de nomes
nomes_personalizados = [
    "GENOVEVA", "ZITA", "ALMA", "HERTA", "FAUSTA", "EURIDICE", "ZELI",
    "ILIDIA", "OTAVIANA", "QUIRINO", "FLORISA", "CASILDA", "SERAFIN",
    "CELSA", "ONDINA", "VICENTINA", "LAUDELINA", "BENVINDA", "DIONIZIA",
    "HELMA", "MERCES", "JUDITI", "EUTALIA", "YEDDA", "CONCEICAO",
    "ADOLPHO", "MIGUELINA", "NEDINA", "IZIDORO", "ALAIDES", "AGUSTINHO",
    "ORIDES", "ILSE", "SIVIRINO", "VITALINA", "OLINDINA", "ZELINDA",
    "GENY", "ZELINA", "PERCILIA", "ZEFERINO", "LIBERATO", "RAYMUNDO",
    "NAZIRA", "EMIDIA", "VERGILIO", "VENANCIA", "GODOFREDO", "LEONTINA",
    "ALBINA", "LEOCADIA", "FLORIZA", "JARDELINA", "SILVERIA", "GUILHERMINO",
    "CIDALIA", "AMANDIO", "HOLANDA", "SATURNINO", "CENIRA", "VALDEMIRA",
    "OLINDO", "CLEMENTINO", "ETELVINO", "JUVENTINO", "AMELIO", "ATALIBA",
    "ANILDA", "FILOMENO", "DELZA", "LEONIDIO", "ZULMIRA", "ERMELINDA",
    "BALBINA", "FIRMINA", "ANTONINA", "JANUARIA", "GERALDINO", "BELARMINO",
    "OLINTO", "ENEDINO", "CASIMIRO", "DEOLINDA", "MALVINA", "AVELINA",
    "OFELIA", "ISALRA", "ESTANISLAU", "IDALINO", "CACILDA", "ITALIA",
    "GERCINA", "OZORIO", "GUILERMINA", "TEODOMIRO", "ERNESTINA", "ANESIA",
    "DELMIRA", "ABADIA", "BENIGNO", "EUPIDIO", "OTACILIA", "ORLANDA",
    "VITALINO", "BERNADINO", "TERTULIANO", "ORESTE", "LEVINO", "FLORINDA",
    "CELESTINA", "ERONDINA", "HERCILIA", "EVANGELINA", "CELITA", "JOVENTINO",
    "ROSITA", "ALCINA", "AMBROSIO", "MARCOLINO", "VALDIVINA", "GERALDINA",
    "CARMELINA", "AMERICA", "ADALGIZA", "JULITA", "GUMERCINDO", "GENESIA",
    "OLIVA", "NARCISA", "ORESTES", "EPIFANIO", "HAYDEE", "RUFINO", "ROSALVA",
    "LAURENTINO", "AURELINA", "IGNES", "CELESTINO", "MODESTO", "INOCENCIO",
    "DELFINO", "EMA", "IRACY", "AURELINO", "OSCARINA", "GONCALA", "CORNELIO",
    "FRANCINA", "DILCE", "JOVINA", "AUZIRA", "EMMA", "ZILA", "ADILIA",
    "SANTINO", "VICENCIA", "DJANIRA", "ISIDORO", "CLARINDO", "DOROTEIA",
    "ZORAIDE", "LEONIDA", "ONORIO", "AGNELO", "HERMINIA", "ALTINA",
    "ARGEMIRO", "EDITH", "SANTO", "PERPETUA", "ALMERINDO", "CECI",
    "SANTINA", "CLOTILDE", "ALTAMIRA", "WALDOMIRO", "PIEDADE", "SILVINA",
    "DORACI", "JURACY", "CARMINA", "ESTELINA", "FORTUNATO", "JESUINA",
    "SEVERIANO", "AGRIPINO", "ZULEICA", "CONSTANTINO", "BERTA", "URBANO",
    "ETELVINA", "ODILIA", "CARLINDA", "DAVINA", "DARCY", "IRENA", "MOACYR",
    "PALMIRA", "HONORIO", "ALVINA", "JUSTINA", "JOVITA", "FRANCELINA",
    "CARMELIA", "CREMILDA", "FLORENCIO", "OSORIO", "JOVELINO", "ARMINDA",
    "DIONISIA", "DINORA", "ANIZIA", "CARMELA", "CLAUDINA", "FELICIANA",
    "CONSUELO", "JOVINO", "OVIDIO", "ARACI", "SANTA", "BELMIRO",
    "CONSTANCIA", "JACY", "DALVINA", "BERNARDINO", "NELLY", "SERAFIM",
    "ELPIDIO", "OLIMPIA", "BENTA", "CLELIA", "ANESIO", "JOSINA", "ARMINDO",
    "ERMINIA", "CARLOTA", "ENEDINA", "JUDITH", "LAUDELINO", "GEORGINA",
    "JOSEFINA", "IRMA", "EUVIRA", "MARTINHA", "HERMINIO", "CUSTODIO",
    "ADALGISA", "ZUMIRA", "FRIDA", "ALMIRA", "IDA", "JORGINA", "CLARINDA",
    "JOVELINA", "LAURITA", "MERCEDES", "CLEMENTE", "GUIOMAR", "OTILIA",
    "EFIGENIA", "DAGMAR", "EULINA", "DEJANIRA", "LEONILDA", "LUCINDA",
    "DIRCE", "WALDEMAR", "ANISIA", "DOLORES", "THEREZA", "ESTELITA",
    "ELSA", "FILOMENA", "EURIDES", "ERCILIA", "CORINA", "ONOFRE",
    "ALMERINDA", "JUREMA", "GENTIL", "LAURINDA", "ARLINDA", "LINDAURA",
    "GENI", "JOAQUINA", "LEONOR", "TEODORA", "JACI", "AVELINO",
    "OLINDA", "ALBINO", "MARIETA", "ARISTIDES", "OLIMPIO", "ANTONIETA",
    "ADELINA", "GERALDA", "ALBERTINA", "MATILDE", "ALEXANDRINA",
    "LAZARA", "IDALINA", "NADIR", "SEVERINA", "ELVIRA", "TERESINHA",
    "GUILHERMINA", "INACIA", "AUGUSTA", "ODETE", "CARMELITA", "OSWALDO",
    "EDITE", "DIVA", "JULIETA", "IRACI", "IZAURA", "ALZIRA", "ZILDA",
    "AURORA", "JUDITE", "NORMA", "CANDIDA", "JANDIRA", "NEUZA",
    "ROSALINA", "NILZA", "ALAIDE", "ADELAIDE", "HILDA", "IRACEMA", "ILDA",
    "NAIR", "MARGARIDA", "ISAURA", "ELZA", "TEREZINHA", "BENEDITA",
    "AMELIA", "SEBASTIANA", "OLGA", "CONCEICAO", "IRENE", "RAIMUNDA",
    "Jose", "Joao", "Antonio", "Francisco", "Carlos", "Paulo", "Pedro",
    "Lucas", "Luiz", "Marcos", "Luis", "Gabriel", "Rafael", "Daniel",
    "Marcelo", "Bruno", "Eduardo", "Felipe", "Raimundo", "Rodrigo",
    "Maria", "Jose", "Ana", "Joao", "Antonio", "Francisco", "Carlos",
    "Paulo", "Pedro", "Lucas", "Luiz", "Marcos", "Luis", "Gabriel",
    "Rafael", "Francisca", "Daniel", "Marcelo", "Bruno", "Eduardo",
]

# Cria uma lista de padrões para o EntityRuler
patterns = [{"label": "PESSOA", "pattern": name} for name in set(nomes_personalizados)]

# --- MODELO PERSONALIZADO ---
# Carrega o modelo base do spaCy em português
try:
    nlp_custom = spacy.load("pt_core_news_sm")
except IOError:
    print("O modelo 'pt_core_news_sm' não foi encontrado. Por favor, execute no seu terminal:")
    print("python -m spacy download pt_core_news_sm")
    exit()

# Cria e adiciona o EntityRuler ao pipeline
ruler = EntityRuler(nlp_custom, overwrite_ents=True)
ruler.add_patterns(patterns)
nlp_custom.add_pipe(ruler)

# --- CÓDIGO DA APLICAÇÃO (views.py e main.py combinados) ---
app = Flask(__name__)
app.secret_key = "segredo-muito-seguro"

# Exemplo de como usar o modelo personalizado
def processar_texto_com_nomes(texto_para_analise):
    """
    Processa um texto e retorna uma lista de nomes próprios de pessoas.
    """
    doc = nlp_custom(texto_para_analise)
    nomes_encontrados = [ent.text for ent in doc.ents if ent.label_ == "PESSOA"]
    return nomes_encontrados

@app.route("/",  methods=["GET", "POST"])
def homepage():
    return "Servidor rodando! Use outra rota para processar texto."

@app.route('/analisar_nomes', methods=['POST'])
def analisar_nomes_rota():
    texto_do_formulario = request.form.get("texto_analise")
    if texto_do_formulario:
        nomes = processar_texto_com_nomes(texto_do_formulario)
        return jsonify({"nomes_encontrados": nomes})
    return jsonify({"nomes_encontrados": []})

nlp_custom = spacy.load("./modelo_nomes_proprios")

def processar_texto_com_nomes(texto_para_analise):
    """
    Processa um texto e retorna uma lista de nomes próprios de pessoas.
    """
    doc = nlp_custom(texto_para_analise)
    nomes_encontrados = [ent.text for ent in doc.ents if ent.label_ == "PESSOA"]
    return nomes_encontrados

@app.route('/analisar_nomes', methods=['POST'])
def analisar_nomes_rota():
    texto_do_formulario = request.form.get("texto_analise")
    if texto_do_formulario:
        nomes = processar_texto_com_nomes(texto_do_formulario)
        return jsonify({"nomes_encontrados": nomes})
    return jsonify({"nomes_encontrados": []})

from regex_patterns import PADROES_SENSIVEIS

# Carrega o modelo do spaCy uma única vez para otimizar o desempenho.
try:
    nlp = spacy.load("pt_core_news_md")
except OSError:
    print("Modelo 'pt_core_news_md' não encontrado. Por favor, baixe-o com 'python -m spacy download pt_core_news_md'")
    nlp = None

app.secret_key = "segredo-muito-seguro"

@app.route("/", methods=["GET", "POST"])
def homepage():
    return render_template("index.html")

def copiar_e_tarjar(original_doc, padroes):
    novo_doc = Document()
    for par in original_doc.paragraphs:
        texto = par.text
        for nome, regex in padroes.items():
            texto = re.sub(regex, lambda m: "█" * len(m.group()), texto)
        novo_doc.add_paragraph(texto)
    return novo_doc

@app.route('/tarjar_docx', methods=['GET', 'POST'])
def tarjar_docx_preview():
    if request.method == 'POST':
        arquivo = request.files.get("docxfile")
        selecionados = request.form.getlist("itens")

        if not arquivo or not arquivo.filename.endswith('.docx'):
            return "Arquivo inválido. Envie um .docx.", 400

        padroes_ativos = {k: v for k, v in PADROES_SENSIVEIS.items() if k in selecionados}
        conteudo_bytes = arquivo.read()
        file_stream = io.BytesIO(conteudo_bytes)
        doc = Document(file_stream)
        ocorrencias = []
        paragrafos_texto = [par.text for par in doc.paragraphs]
        texto_completo = " ".join(paragrafos_texto)

        # Detecção de nomes com spaCy
        if nlp:
            doc_spacy = nlp(texto_completo)
            char_offset = 0
            for i, p_text in enumerate(paragrafos_texto):
                for entidade in nlp(p_text).ents:
                    if entidade.label_ == "PER":
                        ocorrencias.append({
                            "tipo": "Nome",
                            "texto": entidade.text,
                            "paragrafo": i,
                            "start": entidade.start_char,
                            "end": entidade.end_char,
                            "id": str(uuid.uuid4())
                        })
                
                # Para evitar duplicatas, não processa nomes do regex se o spaCy estiver ativo.
                if "NOME" in padroes_ativos:
                    del padroes_ativos["NOME"]

        # Detecção com Regex para os outros padrões
        for i, par_text in enumerate(paragrafos_texto):
            for tipo, regex in padroes_ativos.items():
                for m in re.finditer(regex, par_text):
                    ocorrencias.append({
                        "tipo": tipo,
                        "texto": m.group(),
                        "paragrafo": i,
                        "start": m.start(),
                        "end": m.end(),
                        "id": str(uuid.uuid4())
                    })
        
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        with open(temp_path, "wb") as f:
            f.write(conteudo_bytes)
        
        session['doc_ocorrencias'] = ocorrencias
        session['doc_path'] = temp_path

        return render_template("preview_docx.html", ocorrencias=ocorrencias, paragrafos=paragrafos_texto)

    return render_template("tarjar_docx.html", padroes=PADROES_SENSIVEIS.keys())

@app.route("/aplicar_tarjas_docx", methods=["POST"])
def aplicar_tarjas_docx():
    selecionados = request.form.getlist("selecionados")
    trechos_manuais_raw = request.form.get("tarjas_manualmente_adicionadas", "")
    trechos_manuais = [t.strip() for t in trechos_manuais_raw.split("|") if t.strip()]
    ocorrencias = session.get("doc_ocorrencias", [])
    caminho = session.get("doc_path", None)
    if not caminho or not os.path.exists(caminho):
        return "Erro: Arquivo temporário não encontrado.", 400
    doc = Document(caminho)
    paragrafo_edits = {}
    for item in ocorrencias:
        if item["id"] in selecionados:
            idx = item["paragrafo"]
            texto_original = doc.paragraphs[idx].text
            if idx not in paragrafo_edits:
                paragrafo_edits[idx] = texto_original
            start, end = item["start"], item["end"]
            trecho = texto_original[start:end]
            texto_editado = paragrafo_edits[idx].replace(trecho, "█" * len(trecho), 1)
            paragrafo_edits[idx] = texto_editado
    if trechos_manuais:
        for i, par in enumerate(doc.paragraphs):
            texto = paragrafo_edits.get(i, par.text)
            for trecho_manual in trechos_manuais:
                if trecho_manual in texto:
                    texto = texto.replace(trecho_manual, "█" * len(trecho_manual))
                    paragrafo_edits[i] = texto
    for i, novo_texto in paragrafo_edits.items():
        par = doc.paragraphs[i]
        par.clear()
        run = par.add_run(novo_texto)
        run.font.color.rgb = RGBColor(0, 0, 0)
    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)
    return send_file(
        mem_file,
        as_attachment=True,
        download_name="documento_tarjado.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/atualizar_preview_docx", methods=["POST"])
def atualizar_preview_docx():
    try:
        data = request.get_json(force=True)
        selecionados = set(data.get("selecionados", []))
        trechos_manuais = data.get("manuais", [])
        ocorrencias = session.get("doc_ocorrencias", [])
        caminho = session.get("doc_path", None)
        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400
        doc = Document(caminho)
        paragrafos_atualizados = []
        paragrafos_texto = [par.text for par in doc.paragraphs]
        for i, texto_paragrafo in enumerate(paragrafos_texto):
            for item in sorted(ocorrencias, key=lambda x: x['start'], reverse=True):
                if item["paragrafo"] == i and item["id"] in selecionados:
                    start, end = item["start"], item["end"]
                    texto_paragrafo = texto_paragrafo[:start] + f"<span class='tarja-preview'>" + texto_paragrafo[start:end] + "</span>" + texto_paragrafo[end:]
            for trecho_manual in trechos_manuais:
                def tarja_substituicao(match):
                    return f"<span class='tarja-preview'>{match.group(0)}</span>"
                texto_paragrafo = re.sub(re.escape(trecho_manual), tarja_substituicao, texto_paragrafo, flags=re.IGNORECASE)
            paragrafos_atualizados.append(texto_paragrafo)
        return jsonify({"paragrafos": paragrafos_atualizados})
    except Exception as e:
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

@app.route('/tarjar_pdf', methods=['GET', 'POST'])
def tarjar_pdf():
    if request.method == 'POST':
        arquivo = request.files.get('pdffile')
        tipos_selecionados = request.form.getlist('tipos')
        if not arquivo or not arquivo.filename.endswith('.pdf'):
            return "Arquivo inválido. Envie um .pdf.", 400
        padroes_filtrados = {k: v for k, v in PADROES_SENSIVEIS.items() if k in tipos_selecionados}
        pdf_bytes = arquivo.read()
        doc = fitz.open("pdf", pdf_bytes)
        ocorrencias = []
        redactions_por_pagina = {}
        for pagina_num in range(len(doc)):
            pagina = doc[pagina_num]
            texto = pagina.get_text("text")
            for tipo, regex in padroes_filtrados.items():
                for m in re.finditer(regex, texto):
                    termo = m.group()
                    ocorrencias.append({
                        "id": f"{pagina_num}_{m.start()}_{m.end()}",
                        "tipo": tipo,
                        "texto": termo,
                        "pagina": pagina_num,
                        "start": m.start(),
                        "end": m.end()
                    })
                    areas = pagina.search_for(termo)
                    for area in areas:
                        redactions_por_pagina.setdefault(pagina_num, []).append(area)
        for pagina_idx, areas in redactions_por_pagina.items():
            pagina = doc[pagina_idx]
            for area in areas:
                pagina.add_redact_annot(area, fill=(0, 0, 0))
            pagina.apply_redactions()
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        doc.close()
        pdf_b64 = base64.b64encode(mem_file.read()).decode('utf-8')
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        with open(temp_path, 'wb') as f:
            f.write(pdf_bytes)
        session['pdf_path'] = temp_path
        session['pdf_ocorrencias'] = ocorrencias
        return render_template("preview_pdf.html", ocorrencias=ocorrencias, pdf_data=pdf_b64)
    return render_template('tarjar_pdf.html', padroes=PADROES_SENSIVEIS.keys())

@app.route('/aplicar_tarjas_pdf', methods=['POST'])
def aplicar_tarjas_pdf():
    selecionados = request.form.getlist('selecionados')
    preservar_logo = request.form.get('preservar_logo', '0') == '1'
    trechos_manuais_raw = request.form.get('tarjas_manualmente_adicionadas', '')
    trechos_manuais = [t.strip() for t in trechos_manuais_raw.split('|') if t.strip()]
    caminho = session.get('pdf_path')
    ocorrencias = session.get('pdf_ocorrencias', [])
    if not caminho or not os.path.exists(caminho):
        return "Erro: arquivo temporário não encontrado.", 400
    doc = fitz.open(caminho)
    redactions_por_pagina = {}
    for item in ocorrencias:
        if item['id'] in selecionados:
            pagina_idx = item['pagina']
            termo = item['texto']
            pagina = doc[pagina_idx]
            areas = pagina.search_for(termo)
            for area in areas:
                if preservar_logo and area.y0 < 100:
                    continue
                redactions_por_pagina.setdefault(pagina_idx, []).append(area)
    if trechos_manuais:
        for num_pagina in range(len(doc)):
            pagina = doc[num_pagina]
            texto_pagina = pagina.get_text()
            for trecho in trechos_manuais:
                if trecho in texto_pagina:
                    areas = pagina.search_for(trecho)
                    for area in areas:
                        if preservar_logo and area.y0 < 100:
                            continue
                        redactions_por_pagina.setdefault(num_pagina, []).append(area)
    for pagina_idx, areas in redactions_por_pagina.items():
        pagina = doc[pagina_idx]
        for area in areas:
            pagina.add_redact_annot(area, fill=(0, 0, 0))
        pagina.apply_redactions()
    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)
    doc.close()
    return send_file(
        mem_file,
        as_attachment=True,
        download_name="documento_tarjado.pdf",
        mimetype="application/pdf"
    )

@app.route('/preview_pdf', methods=['POST'])
def preview_pdf():
    arquivo = request.files['arquivo']
    nome_temporario = os.path.join('uploads', f"{uuid.uuid4()}.pdf")
    arquivo.save(nome_temporario)
    doc = fitz.open(nome_temporario)
    pdf_data = base64.b64encode(open(nome_temporario, "rb").read()).decode('utf-8')
    ocorrencias = detectar_dados(doc)
    texto_extraido = ""
    for pagina in doc:
        texto_extraido += pagina.get_text() + "\n"
    doc.close()
    session['pdf_path'] = nome_temporario
    session['pdf_ocorrencias'] = ocorrencias
    return render_template(
        "preview_pdf.html",
        pdf_data=pdf_data,
        ocorrencias=ocorrencias,
        texto_extraido=texto_extraido
    )

@app.route('/atualizar_preview_pdf', methods=['POST'])
def atualizar_preview_pdf():
    try:
        data = request.get_json(force=True)
        selecionados = data.get("selecionados", [])
        trechos_manuais = data.get("manuais", [])
        caminho = session.get('pdf_path')
        ocorrencias = session.get('pdf_ocorrencias', [])
        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400
        doc = fitz.open(caminho)
        redactions_por_pagina = {}
        for item in ocorrencias:
            if item['id'] in selecionados:
                pagina_idx = item['pagina']
                termo = item['texto']
                pagina = doc[pagina_idx]
                areas = pagina.search_for(termo)
                for area in areas:
                    redactions_por_pagina.setdefault(pagina_idx, []).append(area)
        for num_pagina in range(len(doc)):
            pagina = doc[num_pagina]
            texto_pagina = pagina.get_text()
            for trecho in trechos_manuais:
                if trecho in texto_pagina:
                    areas = pagina.search_for(trecho)
                    for area in areas:
                        redactions_por_pagina.setdefault(num_pagina, []).append(area)
        for pagina_idx, areas in redactions_por_pagina.items():
            pagina = doc[pagina_idx]
            for area in areas:
                pagina.add_redact_annot(area, fill=(0, 0, 0))
            pagina.apply_redactions()
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        doc.close()
        pdf_b64 = base64.b64encode(mem_file.read()).decode('utf-8')
        return jsonify({"pdf_data": pdf_b64})
    except Exception as e:
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

@app.route('/download_pdf_tarjado')
def download_pdf_tarjado():
    path = session.get('pdf_tarjado_path', None)
    if not path or not os.path.exists(path):
        return "Nenhum PDF tarjado disponível.", 400
    return send_file(path, as_attachment=True, download_name="documento_tarjado.pdf", mimetype="application/pdf")

@app.route('/tarjar_ocr_pdf', methods=['GET', 'POST'])
def tarjar_ocr_pdf():
    if request.method == 'POST':
        arquivo = request.files.get('ocrpdf')
        tipos_selecionados = request.form.getlist('tipos')
        if not arquivo or not arquivo.filename.lower().endswith('.pdf'):
            return "Arquivo inválido. Envie um arquivo PDF escaneado.", 400
        padroes_ativos = {k: v for k, v in PADROES_SENSIVEIS.items() if k in tipos_selecionados}
        try:
            pdf_bytes = arquivo.read()
            imagens = convert_from_bytes(pdf_bytes)
        except Exception as e:
            app.logger.error(f"Erro ao converter PDF em imagens: {e}")
            return "Erro ao processar o arquivo PDF.", 500
        todas_ocorrencias = []
        for idx, imagem in enumerate(imagens):
            dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)
            for tipo, regex in padroes_ativos.items():
                try:
                    pattern = regex if isinstance(regex, re.Pattern) else re.compile(regex, re.IGNORECASE | re.UNICODE)
                except re.error as e:
                    app.logger.error(f"Regex inválido para tipo '{tipo}': {e}")
                    continue
                for i, palavra in enumerate(dados_ocr['text']):
                    texto = (palavra or '').strip()
                    if not texto:
                        continue
                    if pattern.search(texto):
                        todas_ocorrencias.append({
                            "id": str(uuid.uuid4()),
                            "pagina": idx,
                            "tipo": tipo,
                            "texto": texto
                        })
        texto_manual = request.form.get('tarjas_manualmente_adicionadas', '').strip()
        if texto_manual:
            trechos_manualmente_adicionados = [t.strip() for t in texto_manual.split('|') if t.strip()]
            for idx, imagem in enumerate(imagens):
                dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)
                for trecho in trechos_manualmente_adicionados:
                    todas_ocorrencias.append({
                        "id": str(uuid.uuid4()),
                        "pagina": idx,
                        "tipo": "manual",
                        "texto": trecho
                    })
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        with open(temp_path, 'wb') as f:
            f.write(pdf_bytes)
        session['ocr_original_pdf_path'] = temp_path
        session['ocr_ocorrencias'] = todas_ocorrencias
        app.logger.info(f"Arquivo OCR salvo em: {temp_path}")
        return render_template(
            "preview_ocr.html",
            ocorrencias=todas_ocorrencias
        )
    return render_template('tarjar_ocr_pdf.html', padroes=PADROES_SENSIVEIS.keys())

@app.route('/aplicar_tarjas_ocr_pdf', methods=['POST'])
def aplicar_tarjas_ocr_pdf():
    caminho = session.get('ocr_original_pdf_path')
    ocorrencias_automaticas = session.get('ocr_ocorrencias', [])
    if not caminho or not os.path.exists(caminho):
        return "Arquivo OCR não encontrado.", 400
    imagens = convert_from_bytes(open(caminho, 'rb').read())
    imagens_tarjadas = [img.copy() for img in imagens]
    selecionados = request.form.getlist('selecionados')
    selecionados_set = set(str(s) for s in selecionados)
    texto_manual = request.form.get('tarjas_manuais_json', '').strip()
    trechos_manuais = [t.strip().lower() for t in texto_manual.split('|') if t.strip()]
    tarjas_aplicadas = []
    for idx, imagem in enumerate(imagens_tarjadas):
        draw = ImageDraw.Draw(imagem)
        dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)
        linhas = {}
        for i in range(len(dados_ocr['text'])):
            linha_num = dados_ocr['line_num'][i]
            if linha_num not in linhas:
                linhas[linha_num] = []
            linhas[linha_num].append({
                'text': (dados_ocr['text'][i] or '').strip(),
                'left': int(dados_ocr['left'][i]),
                'top': int(dados_ocr['top'][i]),
                'width': int(dados_ocr['width'][i]),
                'height': int(dados_ocr['height'][i])
            })
        for ocorrencia in [o for o in ocorrencias_automaticas if o['pagina'] == idx]:
            if str(ocorrencia['id']) not in selecionados_set:
                continue
            termo_lower = ocorrencia['texto'].lower()
            for palavras_linha in linhas.values():
                linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                if termo_lower in linha_texto:
                    char_count = 0
                    trecho_start = linha_texto.find(termo_lower)
                    trecho_end = trecho_start + len(termo_lower)
                    for palavra in palavras_linha:
                        palavra_start = char_count
                        palavra_end = char_count + len(palavra['text'])
                        char_count += len(palavra['text']) + 1
                        if palavra_end > trecho_start and palavra_start < trecho_end:
                            x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                            draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                            tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})
        for trecho in trechos_manuais:
            trecho_lower = trecho.lower()
            for palavras_linha in linhas.values():
                linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                if fuzz.partial_ratio(trecho_lower, linha_texto) >= 85:
                    char_count = 0
                    trecho_start = linha_texto.find(trecho_lower)
                    trecho_end = trecho_start + len(trecho_lower)
                    for palavra in palavras_linha:
                        palavra_start = char_count
                        palavra_end = char_count + len(palavra['text'])
                        char_count += len(palavra['text']) + 1
                        if palavra_end > trecho_start and palavra_start < trecho_end:
                            x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                            draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                            tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})
    session['tarjas_ocr'] = tarjas_aplicadas
    buffer = io.BytesIO()
    imagens_tarjadas[0].save(buffer, format="PDF", save_all=True, append_images=imagens_tarjadas[1:])
    buffer.seek(0)
    app.logger.info("PDF gerado com sucesso, enviando para download...")
    return send_file(
        buffer,
        as_attachment=True,
        download_name="documento_tarjado.pdf",
        mimetype="application/pdf"
    )

@app.route('/atualizar_preview_ocr_pdf', methods=['POST'])
def atualizar_preview_ocr_pdf():
    try:
        data = request.get_json(force=True)
        selecionados = data.get("selecionados", [])
        manuais = data.get("manuais", [])
        selecionados_set = set(str(s) for s in selecionados)
        caminho = session.get('ocr_original_pdf_path')
        ocorrencias = session.get('ocr_ocorrencias', [])
        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400
        imagens = convert_from_bytes(open(caminho, 'rb').read())
        imagens_tarjadas = [img.copy() for img in imagens]
        tarjas_aplicadas = []
        for idx, imagem in enumerate(imagens_tarjadas):
            draw = ImageDraw.Draw(imagem)
            dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)
            linhas = {}
            for i in range(len(dados_ocr['text'])):
                linha_num = dados_ocr['line_num'][i]
                if linha_num not in linhas:
                    linhas[linha_num] = []
                linhas[linha_num].append({
                    'text': (dados_ocr['text'][i] or '').strip(),
                    'left': int(dados_ocr['left'][i]),
                    'top': int(dados_ocr['top'][i]),
                    'width': int(dados_ocr['width'][i]),
                    'height': int(dados_ocr['height'][i])
                })
            for ocorrencia in [o for o in ocorrencias if o['pagina'] == idx]:
                if str(ocorrencia['id']) not in selecionados_set:
                    continue
                termo_lower = ocorrencia['texto'].lower()
                for palavras_linha in linhas.values():
                    linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                    if termo_lower in linha_texto:
                        char_count = 0
                        trecho_start = linha_texto.find(termo_lower)
                        trecho_end = trecho_start + len(termo_lower)
                        for palavra in palavras_linha:
                            palavra_start = char_count
                            palavra_end = char_count + len(palavra['text'])
                            char_count += len(palavra['text']) + 1
                            if palavra_end > trecho_start and palavra_start < trecho_end:
                                x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                                draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                                tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})
            for trecho in manuais:
                trecho_lower = trecho.lower()
                for palavras_linha in linhas.values():
                    linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                    if fuzz.partial_ratio(trecho_lower, linha_texto) >= 85:
                        char_count = 0
                        trecho_start = linha_texto.find(trecho_lower)
                        trecho_end = trecho_start + len(trecho_lower)
                        for palavra in palavras_linha:
                            palavra_start = char_count
                            palavra_end = char_count + len(palavra['text'])
                            char_count += len(palavra['text']) + 1
                            if palavra_end > trecho_start and palavra_start < trecho_end:
                                x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                                draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                                tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})
        session['tarjas_ocr'] = tarjas_aplicadas
        pdf_mem = io.BytesIO()
        imagens_tarjadas[0].save(pdf_mem, format="PDF", save_all=True, append_images=imagens_tarjadas[1:])
        pdf_mem.seek(0)
        pdf_b64 = base64.b64encode(pdf_mem.read()).decode('utf-8')
        return jsonify({"pdf_data": pdf_b64})
    except Exception as e:
        app.logger.error(f"Erro ao atualizar preview OCR PDF: {e}")
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

@app.route('/download_pdf_ocr')
def download_pdf_ocr():
    caminho = session.get('ocr_original_pdf_path')
    app.logger.info(f"Caminho do arquivo OCR: {caminho}")
    if not caminho or not os.path.exists(caminho):
        app.logger.error("Arquivo não encontrado.")
        return "Arquivo não encontrado.", 404
    return send_file(caminho, as_attachment=True, download_name="pdf_tarjado_ocr.pdf")

@app.route('/ver_pdf_ocr')
def ver_pdf_ocr():
    caminho = session.get('ocr_original_pdf_path')
    if not caminho or not os.path.exists(caminho):
        app.logger.error("Arquivo não encontrado.")
        return "Arquivo não encontrado.", 404
    return send_file(caminho, mimetype='application/pdf')

if __name__ == "__main__":
    app.run(debug=True)