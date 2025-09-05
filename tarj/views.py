# views.py

from main import app
from flask import render_template, request, send_file, redirect, url_for, session, jsonify
import re
import tempfile
import os
import fitz  # PyMuPDF
from docx import Document
import io
import uuid
import base64
from docx.shared import RGBColor
from fuzzywuzzy import fuzz
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image, ImageDraw

# A importação agora vai funcionar corretamente
from historico_utils import salvar_envio
from regex_patterns import PADROES_SENSIVEIS

# Habilita sessão para guardar dados temporários
app.secret_key = "coloque-sua-chave-secreta-aqui-e-mude-para-producao"


@app.route("/", methods=["GET"])
def homepage():
    return render_template("index.html")

# -----------------------------------------------------------------------------------
# FUNÇÕES PARA DOCX
# -----------------------------------------------------------------------------------

@app.route('/tarjar_docx', methods=['GET', 'POST'])
def tarjar_docx_preview():
    if request.method == 'POST':
        arquivo = request.files.get("docxfile")
        selecionados = request.form.getlist("itens")

        if not arquivo or not arquivo.filename.endswith('.docx'):
            return "Arquivo inválido. Envie um .docx.", 400

        padroes_ativos = {k: v for k, v in PADROES_SENSIVEIS.items() if k in selecionados}

        conteudo_bytes = arquivo.read()
        file_stream = io.BytesIO(conteudo_bytes)
        doc = Document(file_stream)

        ocorrencias = []
        paragrafos_com_tarja = []

        for i, par in enumerate(doc.paragraphs):
            texto_original = par.text
            texto_tarjado = texto_original
            
            for tipo, regex in padroes_ativos.items():
                for m in re.finditer(regex, texto_original):
                    encontrado = m.group()
                    texto_tarjado = texto_tarjado.replace(encontrado, '█' * len(encontrado))

                    ocorrencias.append({
                        "tipo": tipo,
                        "texto": encontrado,
                        "paragrafo": i,
                        "start": m.start(),
                        "end": m.end(),
                        "id": f"{i}_{m.start()}_{m.end()}"
                    })
            paragrafos_com_tarja.append(texto_tarjado)

        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        with open(temp_path, "wb") as f:
            f.write(conteudo_bytes)

        session['doc_ocorrencias'] = ocorrencias
        session['doc_path'] = temp_path
        session['original_filename'] = arquivo.filename

        return render_template("preview_docx.html", ocorrencias=ocorrencias, paragrafos=paragrafos_com_tarja)

    return render_template("tarjar_docx.html", padroes=PADROES_SENSIVEIS.keys())

@app.route("/aplicar_tarjas_docx", methods=["POST"])
def aplicar_tarjas_docx():
    selecionados = request.form.getlist("selecionados")
    trechos_manuais_raw = request.form.get("tarjas_manualmente_adicionadas", "")
    trechos_manuais = [t.strip() for t in trechos_manuais_raw.split("|") if t.strip()]

    ocorrencias = session.get("doc_ocorrencias", [])
    caminho = session.get("doc_path", None)
    filename = session.get("original_filename", "documento.docx")

    if not caminho or not os.path.exists(caminho):
        return "Erro: Arquivo temporário não encontrado.", 400

    doc = Document(caminho)
    paragrafos_para_editar = {}

    for item in ocorrencias:
        if item["id"] in selecionados:
            idx = item["paragrafo"]
            if idx not in paragrafos_para_editar:
                paragrafos_para_editar[idx] = list(doc.paragraphs[idx].text)
            
            start, end = item["start"], item["end"]
            for i in range(start, end):
                if i < len(paragrafos_para_editar[idx]):
                    paragrafos_para_editar[idx][i] = '█'

    textos_editados = {idx: "".join(chars) for idx, chars in paragrafos_para_editar.items()}

    for i, par in enumerate(doc.paragraphs):
        texto_atual = textos_editados.get(i, par.text)
        for trecho in trechos_manuais:
            texto_atual = texto_atual.replace(trecho, "█" * len(trecho))
        textos_editados[i] = texto_atual

    for i, novo_texto in textos_editados.items():
        par = doc.paragraphs[i]
        par.clear()
        par.add_run(novo_texto)

    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)

    session.pop('doc_ocorrencias', None)
    session.pop('doc_path', None)
    session.pop('original_filename', None)
    
    salvar_envio(filename, 'docx')

    return send_file(
        mem_file,
        as_attachment=True,
        download_name="documento_tarjado.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/atualizar_preview_docx", methods=["POST"])
def atualizar_preview_docx():
    try:
        data = request.get_json(force=True)
        selecionados = set(data.get("selecionados", []))
        trechos_manuais = data.get("manuais", [])

        ocorrencias = session.get("doc_ocorrencias", [])
        caminho = session.get("doc_path", None)

        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400

        doc = Document(caminho)
        paragrafos_para_editar = {}

        for item in ocorrencias:
            if item["id"] in selecionados:
                idx = item["paragrafo"]
                if idx not in paragrafos_para_editar:
                    paragrafos_para_editar[idx] = list(doc.paragraphs[idx].text)
                start, end = item["start"], item["end"]
                for i in range(start, end):
                    if i < len(paragrafos_para_editar[idx]):
                        paragrafos_para_editar[idx][i] = '█'
        
        textos_editados = {idx: "".join(chars) for idx, chars in paragrafos_para_editar.items()}

        paragrafos_atualizados = []
        for i, par in enumerate(doc.paragraphs):
            texto_atual = textos_editados.get(i, par.text)
            for trecho in trechos_manuais:
                texto_atual = re.sub(re.escape(trecho), "█" * len(trecho), texto_atual, flags=re.IGNORECASE)
            paragrafos_atualizados.append(texto_atual)

        return jsonify({"paragrafos": paragrafos_atualizados})

    except Exception as e:
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

# -----------------------------------------------------------------------------------
# FUNÇÕES PARA PDF (TEXTO)
# -----------------------------------------------------------------------------------

@app.route('/tarjar_pdf', methods=['GET', 'POST'])
def tarjar_pdf():
    if request.method == 'POST':
        arquivo = request.files.get('pdffile')
        tipos_selecionados = request.form.getlist('tipos')  

        if not arquivo or not arquivo.filename.endswith('.pdf'):
            return "Arquivo inválido. Envie um .pdf.", 400

        padroes_filtrados = {k: v for k, v in PADROES_SENSIVEIS.items() if k in tipos_selecionados}
        pdf_bytes = arquivo.read()
        
        ocorrencias = []
        with fitz.open("pdf", pdf_bytes) as doc:
            for pagina_num, pagina in enumerate(doc):
                texto = pagina.get_text("text")
                for tipo, regex in padroes_filtrados.items():
                    for m in re.finditer(regex, texto):
                        ocorrencias.append({
                            "id": f"{pagina_num}_{m.start()}_{m.end()}",
                            "tipo": tipo,
                            "texto": m.group(),
                            "pagina": pagina_num
                        })
            
            # Cria PDF de preview
            with fitz.open("pdf", pdf_bytes) as doc_preview:
                for item in ocorrencias:
                    pagina = doc_preview[item['pagina']]
                    areas = pagina.search_for(item['texto'])
                    for area in areas:
                        pagina.add_redact_annot(area, fill=(0, 0, 0))
                
                for pagina in doc_preview:
                    pagina.apply_redactions()

                mem_file = io.BytesIO()
                doc_preview.save(mem_file)
                mem_file.seek(0)
                pdf_b64 = base64.b64encode(mem_file.read()).decode('utf-8')

        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        with open(temp_path, 'wb') as f:
            f.write(pdf_bytes)

        session['pdf_path'] = temp_path
        session['pdf_ocorrencias'] = ocorrencias
        session['original_filename'] = arquivo.filename

        return render_template("preview_pdf.html", ocorrencias=ocorrencias, pdf_data=pdf_b64)

    return render_template('tarjar_pdf.html', padroes=PADROES_SENSIVEIS.keys())

@app.route('/aplicar_tarjas_pdf', methods=['POST'])
def aplicar_tarjas_pdf():
    selecionados = request.form.getlist('selecionados')
    trechos_manuais_raw = request.form.get('tarjas_manualmente_adicionadas', '')
    trechos_manuais = [t.strip() for t in trechos_manuais_raw.split('|') if t.strip()]

    caminho = session.get('pdf_path')
    ocorrencias = session.get('pdf_ocorrencias', [])
    filename = session.get("original_filename", "documento.pdf")

    if not caminho or not os.path.exists(caminho):
        return "Erro: arquivo temporário não encontrado.", 400

    with fitz.open(caminho) as doc:
        # Tarjas automáticas
        for item in ocorrencias:
            if item['id'] in selecionados:
                pagina = doc[item['pagina']]
                areas = pagina.search_for(item['texto'])
                for area in areas:
                    pagina.add_redact_annot(area, fill=(0, 0, 0))
        
        # Tarjas manuais
        if trechos_manuais:
            for pagina in doc:
                for trecho in trechos_manuais:
                    areas = pagina.search_for(trecho)
                    for area in areas:
                        pagina.add_redact_annot(area, fill=(0, 0, 0))

        # Aplica todas as tarjas de uma vez
        for pagina in doc:
            pagina.apply_redactions()
        
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)

    # Limpa sessão
    session.pop('pdf_path', None)
    session.pop('pdf_ocorrencias', None)
    session.pop('original_filename', None)

    # Salva no histórico
    salvar_envio(filename, 'pdf')

    return send_file(
        mem_file,
        as_attachment=True,
        download_name="documento_tarjado.pdf",
        mimetype="application/pdf"
    )

@app.route('/atualizar_preview_pdf', methods=['POST'])
def atualizar_preview_pdf():
    try:
        data = request.get_json(force=True)
        selecionados = data.get("selecionados", [])
        trechos_manuais = data.get("manuais", [])

        caminho = session.get('pdf_path')
        ocorrencias = session.get('pdf_ocorrencias', [])

        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400

        with fitz.open(caminho) as doc:
            # Tarjas automáticas
            for item in ocorrencias:
                if item['id'] in selecionados:
                    pagina = doc[item['pagina']]
                    areas = pagina.search_for(item['texto'])
                    for area in areas:
                        pagina.add_redact_annot(area, fill=(0, 0, 0))
            
            # Tarjas manuais
            if trechos_manuais:
                for pagina in doc:
                    for trecho in trechos_manuais:
                        areas = pagina.search_for(trecho)
                        for area in areas:
                            pagina.add_redact_annot(area, fill=(0, 0, 0))

            for pagina in doc:
                pagina.apply_redactions()

            mem_file = io.BytesIO()
            doc.save(mem_file)
            mem_file.seek(0)
            pdf_b64 = base64.b64encode(mem_file.read()).decode('utf-8')

        return jsonify({"pdf_data": pdf_b64})

    except Exception as e:
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

# -----------------------------------------------------------------------------------
# FUNÇÕES PARA PDF (OCR)
# -----------------------------------------------------------------------------------

def _desenhar_tarjas_em_imagens_ocr(imagens, ocorrencias, selecionados_set, manuais, dados_ocr_por_pagina):
    """Função auxiliar para desenhar tarjas (evita duplicação de código)."""
    imagens_tarjadas = [img.copy() for img in imagens]
    
    for idx, imagem in enumerate(imagens_tarjadas):
        draw = ImageDraw.Draw(imagem)
        dados_ocr = dados_ocr_por_pagina[idx]

        # Processamento de tarjas automáticas selecionadas
        for item in ocorrencias:
            if item['pagina'] == idx and str(item['id']) in selecionados_set:
                termo_busca = item['texto']
                texto_pagina = " ".join(dados_ocr['text'])
                if termo_busca in texto_pagina:
                    for i, palavra in enumerate(dados_ocr['text']):
                        if palavra in termo_busca:
                            (x, y, w, h) = (dados_ocr['left'][i], dados_ocr['top'][i], dados_ocr['width'][i], dados_ocr['height'][i])
                            draw.rectangle([(x, y), (x + w, y + h)], fill='black')
        
        # Processamento de tarjas manuais
        for trecho_manual in manuais:
            texto_pagina_lower = " ".join(p.lower() for p in dados_ocr['text'])
            if fuzz.partial_ratio(trecho_manual.lower(), texto_pagina_lower) > 85:
                for i, palavra in enumerate(dados_ocr['text']):
                    if palavra.lower() in trecho_manual.lower():
                        (x, y, w, h) = (dados_ocr['left'][i], dados_ocr['top'][i], dados_ocr['width'][i], dados_ocr['height'][i])
                        draw.rectangle([(x, y), (x + w, y + h)], fill='black')

    return imagens_tarjadas

@app.route('/tarjar_ocr_pdf', methods=['GET', 'POST'])
def tarjar_ocr_pdf():
    if request.method == 'POST':
        arquivo = request.files.get('ocrpdf')
        tipos_selecionados = request.form.getlist('tipos')

        if not arquivo or not arquivo.filename.lower().endswith('.pdf'):
            return "Arquivo inválido. Envie um PDF.", 400

        padroes_ativos = {k: v for k, v in PADROES_SENSIVEIS.items() if k in tipos_selecionados}
        
        try:
            pdf_bytes = arquivo.read()
            imagens = convert_from_bytes(pdf_bytes)
        except Exception as e:
            return f"Erro ao converter PDF em imagens: {e}", 500

        todas_ocorrencias = []
        dados_ocr_por_pagina = []

        for idx, imagem in enumerate(imagens):
            dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)
            dados_ocr_por_pagina.append(dados_ocr) # Salva dados completos para reuso

            texto_completo_pagina = " ".join(dados_ocr['text'])
            
            for tipo, regex in padroes_ativos.items():
                for m in re.finditer(regex, texto_completo_pagina):
                    todas_ocorrencias.append({
                        "id": str(uuid.uuid4()),
                        "pagina": idx,
                        "tipo": tipo,
                        "texto": m.group()
                    })
        
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        with open(temp_path, 'wb') as f:
            f.write(pdf_bytes)

        session['ocr_original_pdf_path'] = temp_path
        session['ocr_ocorrencias'] = todas_ocorrencias
        session['ocr_dados_completos'] = dados_ocr_por_pagina # Otimização de performance!
        session['original_filename'] = arquivo.filename

        return render_template("preview_ocr.html", ocorrencias=todas_ocorrencias)

    return render_template('tarjar_ocr_pdf.html', padroes=PADROES_SENSIVEIS.keys())

@app.route('/aplicar_tarjas_ocr_pdf', methods=['POST'])
def aplicar_tarjas_ocr_pdf():
    caminho = session.get('ocr_original_pdf_path')
    ocorrencias = session.get('ocr_ocorrencias', [])
    dados_ocr = session.get('ocr_dados_completos', [])
    filename = session.get("original_filename", "documento_ocr.pdf")

    if not all([caminho, os.path.exists(caminho), dados_ocr]):
        return "Erro: Arquivo OCR ou dados temporários não encontrados.", 400

    selecionados_set = set(request.form.getlist('selecionados'))
    trechos_manuais_raw = request.form.get('tarjas_manualmente_adicionadas', '')
    trechos_manuais = [t.strip() for t in trechos_manuais_raw.split('|') if t.strip()]

    imagens = convert_from_bytes(open(caminho, 'rb').read())
    imagens_tarjadas = _desenhar_tarjas_em_imagens_ocr(
        imagens, ocorrencias, selecionados_set, trechos_manuais, dados_ocr
    )
    
    buffer = io.BytesIO()
    if imagens_tarjadas:
        imagens_tarjadas[0].save(buffer, format="PDF", save_all=True, append_images=imagens_tarjadas[1:])
        buffer.seek(0)

    # Limpa sessão
    session.pop('ocr_original_pdf_path', None)
    session.pop('ocr_ocorrencias', None)
    session.pop('ocr_dados_completos', None)
    session.pop('original_filename', None)

    # Salva no histórico
    salvar_envio(filename, 'pdf_ocr')

    return send_file(
        buffer,
        as_attachment=True,
        download_name="documento_ocr_tarjado.pdf",
        mimetype="application/pdf"
    )

@app.route('/atualizar_preview_ocr_pdf', methods=['POST'])
def atualizar_preview_ocr_pdf():
    try:
        data = request.get_json(force=True)
        selecionados_set = set(data.get("selecionados", []))
        manuais = data.get("manuais", [])

        caminho = session.get('ocr_original_pdf_path')
        ocorrencias = session.get('ocr_ocorrencias', [])
        dados_ocr = session.get('ocr_dados_completos', [])

        if not all([caminho, os.path.exists(caminho), dados_ocr]):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400

        imagens = convert_from_bytes(open(caminho, 'rb').read())
        imagens_tarjadas = _desenhar_tarjas_em_imagens_ocr(
            imagens, ocorrencias, selecionados_set, manuais, dados_ocr
        )

        pdf_mem = io.BytesIO()
        if imagens_tarjadas:
            imagens_tarjadas[0].save(pdf_mem, format="PDF", save_all=True, append_images=imagens_tarjadas[1:])
            pdf_mem.seek(0)
        
        pdf_b64 = base64.b64encode(pdf_mem.read()).decode('utf-8')
        return jsonify({"pdf_data": pdf_b64})

    except Exception as e:
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500